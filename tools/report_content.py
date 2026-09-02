#!/usr/bin/env python3
"""
CloudMatrix - report content
============================

The narrative, tables and evidence layout of the assignment report. Kept apart
from build_report.py so that the document helpers and the document content can
be edited independently.

Run:  python tools/report_content.py
"""

from __future__ import annotations

import json
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

import build_report as B
from build_report import (DOCS, FIG, NAME, NAVY, REGNO, REPO, SHOT, bullet,
                          evidence, h1, h2, h3, image, link, mono, page_numbers,
                          pagebreak, para, setup, table)


# ==========================================================================
# Front matter
# ==========================================================================

def cover(doc):
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("COMMON COURSE ASSIGNMENT")
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = NAVY

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CSA04 — OPERATING SYSTEMS")
    r.bold = True; r.font.size = Pt(13)

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Covering CO3, CO4 and CO5")
    r.font.size = Pt(11); r.font.color.rgb = B.GREY

    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Design, Memory Optimization and System Administration\n"
                  "of an Enterprise Cloud and Virtualization Platform")
    r.bold = True; r.font.size = Pt(17); r.font.color.rgb = NAVY

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("A quantitative study of paging, file allocation, inode dynamics,\n"
                  "disk scheduling and Linux service design for the “CloudMatrix” host")
    r.font.size = Pt(11); r.italic = True; r.font.color.rgb = B.GREY

    doc.add_paragraph()
    image(doc, FIG / "fig08_architecture.png", "", width=6.1, is_figure=False)

    rows = [
        ("Department", "Computer Science and Engineering"),
        ("Programme", "B.E. / B.Tech — CSE"),
        ("Course Code & Name", "CSA04 — Operating Systems"),
        ("Academic Year / Batch", "2026 – 2027"),
        ("Faculty Name", "Dr. Priskilla Angel Rani .J"),
        ("Assignment Title", "Design, Memory Optimization, and System Administration "
                             "of an Enterprise Cloud & Virtualization Platform"),
        ("Course Outcomes", "CO3 — Virtual memory and paging;  CO4 — File systems, "
                            "inode dynamics and disk scheduling;  CO5 — Linux "
                            "administration, network services and virtualization"),
        ("Bloom's Taxonomy Level", "L3 — Apply;  L4 — Analyse;  L5 — Evaluate"),
        ("SDG Mapping", "SDG 7 — Affordable and Clean Energy;  SDG 9 — Industry, "
                        "Innovation and Infrastructure;  SDG 11 — Sustainable Cities"),
        ("Date of Submission", "02 / 09 / 2026"),
        ("Maximum Marks", "100"),
        ("Submitted By", f"{NAME}          Reg. No. {REGNO}"),
        ("Implementation Repository", REPO),
    ]
    B.table(doc, ["Field", "Detail"], rows, "Assignment information.",
            widths=[1.9, 4.6], size=9)


def declaration_and_abstract(doc):
    h1(doc, "Declaration and Abstract")

    h2(doc, "Declaration")
    para(doc,
         "I declare that the work presented in this report is my own. Every numerical "
         "result quoted here was produced by a program written for this assignment "
         "and committed to the public repository listed below; none of the figures "
         "were transcribed by hand or estimated. The console output reproduced in the "
         "evidence sections is the verbatim standard output of those programs, "
         "captured automatically into "
         "results/logs/ by the driver script run_all.py, and the source excerpts "
         "shown are rendered directly from the committed files with their real line "
         "numbers. Any reader may clone the repository and regenerate the entire "
         "evidence base with a single command; if a number in this report and a "
         "number in the repository ever disagree, the repository is correct and the "
         "report is stale.")
    link(doc, REPO, REPO, prefix="Implementation repository:  ")
    para(doc,
         "Where a required target was not met, this report says so and explains why, "
         "rather than presenting a partial result as a success. Section 6.4 contains "
         "one such finding.")

    h2(doc, "Abstract")
    para(doc,
         "**CloudMatrix** is a Linux 6.x enterprise virtualization host carrying 1,200 "
         "concurrent tenant workloads across four tiers — web and DNS microservices, "
         "interactive enterprise applications, database guest virtual machines and "
         "batch analytics workers — on 32 GB of physical memory and a 500-cylinder "
         "block storage unit. This report engineers three subsystems for that host and "
         "defends each design decision with measurements rather than assertions.")
    para(doc,
         "For **memory (CO3)**, the host geometry is derived from first principles: "
         "8,388,608 physical frames and 32,768 virtual pages per process, addressed "
         "through a two-level page table whose 27-bit virtual address splits as "
         "5 | 10 | 12 bits. Sizing the inner table to exactly one frame reduces "
         "resident page-table cost from 128 KB to 4 KB per process, a 32× saving that "
         "matters at 1,200 guests. An executable MMU with a 64-entry TLB measures an "
         "effective access time of 184.33 ns against 300 ns without translation "
         "caching. Page replacement is compared across FIFO, LRU and Optimal on an "
         "18-reference trace, and **Belady's anomaly is confirmed for FIFO** — 11 "
         "faults with three frames rising to 12 with four. That single result decides "
         "the policy, because a balloon driver that continuously resizes guest memory "
         "cannot be built on an algorithm with no monotonicity guarantee.")
    para(doc,
         "For **storage (CO4)**, four allocation strategies are simulated on a shared "
         "device. After a create/delete churn cycle the disk holds 248 free blocks yet "
         "contiguous allocation **refuses** a 150-block request, because the largest "
         "free run is only 95 — external fragmentation measured rather than defined. "
         "A working miniature UNIX file system traces ialloc, ifree, alloc, free and "
         "namei through a guest lifecycle, and shows that a 50 GB disk image costs "
         "12,814 indirect blocks of pure pointer metadata. Six disk-scheduling "
         "algorithms are evaluated on the specified queue, and then again under "
         "continuous arrivals at 88 % device load, where SSTF's attractive 5.2 ms mean "
         "is shown to conceal a 71.2 ms starvation tail.")
    para(doc,
         "For **systems administration (CO5)**, a production-shaped Linux "
         "multifunction server is specified: BIND 9 with TSIG-authenticated transfers, "
         "DNSSEC validation, response rate limiting and GDPR-aware log retention; ISC "
         "DHCP with reservations and a PXE class; bridge networking with a genuinely "
         "isolated tenant overlay; and Xen and libvirt guest definitions with CPU "
         "pinning, ballooning floors set at measured working sets, per-guest I/O "
         "throttling and sVirt labelling.")
    para(doc,
         "The implementation comprises roughly 3,900 lines of Python, shell and "
         "configuration across a public Git repository with a **65-test validation "
         "suite that passes clean**. The tests check both hand-computed values and "
         "invariants that must hold for any correct implementation, so a silent "
         "regression in a simulator is caught rather than published.")


TOC_MAP = Path(__file__).resolve().parent / "toc_pages.json"


def contents(doc):
    h1(doc, "Table of Contents")
    entries = [
        ("Declaration and Abstract", "Declaration and Abstract"),
        ("1.  Problem Understanding and Formulation",
         "1. Problem Understanding and Formulation"),
        ("     1.1  The CloudMatrix problem restated", "1.1  The CloudMatrix problem"),
        ("     1.2  Bottlenecks identified in the scenario", "1.2  Bottlenecks"),
        ("     1.3  Workload assumptions and their justification",
         "1.3  Workload assumptions"),
        ("     1.4  Measurable metrics and success criteria", "1.4  Measurable metrics"),
        ("     1.5  Constraints carried from Section D", "1.5  Constraints carried"),
        ("2.  Application of Course Knowledge", "2. Application of Course Knowledge"),
        ("     2.1  CO3 — Memory layout and two-level page table design",
         "2.1  Memory layout and two-level"),
        ("     2.2  CO3 — Dynamic storage allocation and dynamic linking",
         "2.2  Dynamic storage allocation"),
        ("     2.3  CO3 — Page replacement and Belady's anomaly",
         "2.3  Page replacement and Belady"),
        ("     2.4  CO3 — Working set, page-fault frequency and thrashing",
         "2.4  Demand paging, working set"),
        ("     2.5  CO4 — File allocation strategies", "2.5  File allocation strategies"),
        ("     2.6  CO4 — UNIX inode dynamics and kernel system calls",
         "2.6  UNIX inode dynamics"),
        ("     2.7  CO4 — Disk head movement and scheduling",
         "2.7  Disk head movement"),
        ("     2.8  CO4 — I/O system architecture", "2.8  I/O system architecture"),
        ("     2.9  CO5 — Linux multifunction server configuration",
         "2.9  Linux multifunction server"),
        ("     2.10 CO5 — Hypervisor and virtualization architecture",
         "2.10  Hypervisor and virtualization"),
        ("     2.11 CO5 — Resource isolation and VM management",
         "2.11  Resource isolation"),
        ("3.  Solution, Design and Methodology", "3. Solution, Design and Methodology"),
        ("4.  Use of Modern Tools", "4. Use of Modern Tools"),
        ("5.  Results and Validation", "5. Results and Validation"),
        ("6.  Analysis and Engineering Decisions", "6. Analysis and Engineering"),
        ("7.  Broader Considerations", "7. Broader Considerations"),
        ("8.  Conclusion and Reflection", "8. Conclusion and Reflection"),
        ("9.  References", "9. References"),
        ("Appendix A — Repository map and reproduction", "Appendix A"),
    ]
    pages = {}
    if TOC_MAP.exists():
        pages = json.loads(TOC_MAP.read_text(encoding="utf-8"))
    entries = [(text, pages.get(key, "—")) for text, key in entries]

    for text, pg in entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(text)
        r.font.size = Pt(10)
        if not text.startswith("     "):
            r.bold = True
        tab = p.add_run("  " + "." * max(4, 74 - len(text)) + "  " + str(pg))
        tab.font.size = Pt(9)
        tab.font.color.rgb = B.GREY

    doc.add_paragraph()
    h2(doc, "Note on evidence")
    para(doc,
         "Three kinds of image appear in this report and they are labelled distinctly "
         "so a reader knows exactly what each one is:")
    bullet(doc, "**Figures** are charts and schematics generated by "
                "src/make_figures.py and src/make_diagrams.py directly from the "
                "simulators, so no figure can drift from the data it depicts.")
    bullet(doc, "**Console transcripts** are terminal-framed renderings of the "
                "verbatim standard output captured in results/logs/. The log file "
                "name appears in the title bar of every one, so any line can be "
                "checked against the repository.")
    bullet(doc, "**Source excerpts** are rendered from the committed files with "
                "their real line numbers, so each maps to a permanent link of the "
                "form  .../blob/main/[path]#L[start]-L[end]")


# ==========================================================================
# 1. Problem understanding
# ==========================================================================

def section1(doc):
    h1(doc, "Problem Understanding and Formulation", 1)

    h2(doc, "1.1  The CloudMatrix problem restated")
    para(doc,
         "CloudMatrix is a single physical server that has been asked to behave like "
         "four different machines at once. It runs Linux 6.x with a hypervisor layer "
         "(Xen in production, KVM/libvirt in the lab) and hosts 1,200 client workloads "
         "for public and private sector tenants. Those workloads do not resemble each "
         "other. Web and DNS microservices are numerous, tiny and latency-sensitive. "
         "Interactive enterprise applications are fewer and larger and must feel "
         "responsive to a human being. Database guests are memory-hungry and "
         "random-access. Batch analytics jobs are throughput-oriented and, crucially, "
         "have nobody waiting on them.")
    para(doc,
         "The engineering problem is that these four tiers compete for exactly three "
         "finite resources — memory frames, disk blocks and the disk arm — and the "
         "policy that is best for one tier is frequently worst for another. A "
         "replacement policy tuned for the database guest can starve the web tier. A "
         "disk scheduler that minimises total head movement can leave a tenant's "
         "backup request unserved indefinitely. An allocation strategy that is fastest "
         "for a 50 GB virtual disk image is absurd for a 2 KB configuration file. The "
         "task is therefore not to find the best algorithm in the abstract, but to "
         "quantify the trade-offs precisely enough that a defensible choice can be "
         "made for each tier, and to show what breaks when the wrong one is chosen.")
    para(doc,
         "This report treats that as an empirical question. Rather than reciting the "
         "textbook properties of each algorithm, each one is implemented, run against "
         "a workload derived from the scenario, and measured. Several of the results "
         "were not what the textbook ordering would predict — LRU produces more faults "
         "than FIFO at three frames on this trace, and SSTF and LOOK tie exactly on "
         "the specified queue — and those surprises are reported and explained rather "
         "than smoothed over, because they turn out to carry the most useful "
         "engineering information in the study.")

    h2(doc, "1.2  Bottlenecks identified in the scenario")
    h3(doc, "Memory allocation bottlenecks (CO3)")
    para(doc,
         "The aggregate working set of the tenant mix exceeds physical memory. This is "
         "not a mistake in the workload sizing; it is the normal condition of a cloud "
         "host, which sells memory it does not have on the assumption that not every "
         "guest is hot simultaneously. The measured over-commit ratio for the tier mix "
         "specified in §1.3 is D/m = 1.248, i.e. 24.8 % more demand than supply. Left "
         "alone, this guarantees thrashing. A second, quieter bottleneck is page-table "
         "residency: a flat page table costs 128 KB per process, which across 1,200 "
         "guests is 150 MB of memory holding mostly untouched entries.")

    h3(doc, "File system and block fragmentation (CO4)")
    para(doc,
         "Guests are created and destroyed constantly, so free space becomes "
         "fragmented. The danger is not running out of space but running out of "
         "*contiguous* space, which fails allocations while the free-space counter "
         "still looks healthy. A second bottleneck is metadata reach: the classical "
         "inode addresses a large file through three levels of indirection, and for a "
         "50 GB image that means thousands of pointer blocks that must themselves be "
         "cached and read.")

    h3(doc, "Disk head thrashing (CO4)")
    para(doc,
         "Sequential streaming of virtual disk images and random metadata access "
         "arrive at the same spindle. Greedy scheduling keeps the head in whichever "
         "region is currently busiest, which is efficient on average and catastrophic "
         "for whatever sits at the far edge of the platter.")

    h3(doc, "Network service and tenant isolation risks (CO5)")
    para(doc,
         "A recursive resolver reachable from outside its intended network is not a "
         "misconfiguration, it is an amplification weapon. Unauthenticated zone "
         "transfers hand an attacker a map of the whole estate. A shared layer-2 "
         "segment lets one compromised guest impersonate the gateway. And DNS query "
         "logs contain client IP addresses, which are personal data under GDPR "
         "Article 4 and cannot simply be kept forever.")

    h3(doc, "Hypervisor overhead")
    para(doc,
         "Every abstraction the hypervisor adds costs something. The design question "
         "is which overheads buy isolation worth paying for and which are pure waste — "
         "for example, letting the host page cache duplicate a database guest's own "
         "buffer pool wastes host memory twice over.")

    h2(doc, "1.3  Workload assumptions and their justification")
    para(doc,
         "The brief fixes some parameters and leaves others open. Everything assumed "
         "here is stated explicitly with the reason for the choice, because an "
         "assumption that is not written down cannot be checked.")

    B.table(doc,
            ["Parameter", "Value", "Source / justification"],
            [
                ("Physical RAM", "32 GB", "Fixed by Section D"),
                ("Page / frame size", "4 KB", "Fixed by Section D; 2 MB hugepages "
                                             "evaluated separately in §2.1"),
                ("Logical address space", "128 MB", "Fixed by Section D"),
                ("Page-table entry size", "4 bytes", "Standard 32-bit PTE; makes the "
                                                     "inner table exactly one frame"),
                ("TLB", "64 entries, fully associative",
                 "Typical L1 dTLB size on x86-64 server parts"),
                ("Disk", "500 cylinders, 0–499", "Fixed by Section D"),
                ("Pending I/O queue", "86, 147, 312, 91, 177, 48, 409, 22, 130, 365, "
                                      "220, 480", "Fixed by Section D"),
                ("Initial head position", "cylinder 125, moving upward",
                 "Fixed by Section D"),
                ("Seek model", "0.5 ms settle + 0.01 ms per cylinder",
                 "Mid-range enterprise SAS spindle; applied uniformly so the "
                 "comparison stays fair"),
                ("Page-fault service time", "8 ms",
                 "NVMe-backed swap including queueing"),
                ("Memory access time", "100 ns", "DRAM latency"),
                ("Block size", "4 KB", "Matched to the page size so the page cache and "
                                       "the buffer cache share a unit"),
                ("Pointer size", "4 bytes", "Gives 1,024 pointers per indirect block"),
                ("Inode structure", "10 direct, 1 single, 1 double, 1 triple",
                 "System V / Bach model, as specified in the syllabus"),
                ("Concurrent sessions", "1,200", "Fixed by Section D"),
                ("Management network", "10.20.30.0/24", "Assumed; RFC 1918 private "
                                                        "range"),
                ("Tenant overlay network", "10.20.40.0/24", "Assumed; separate "
                                                            "broadcast domain"),
            ],
            "Assumed and given parameters for the CloudMatrix host.",
            widths=[1.55, 1.9, 3.05], size=7.4)

    para(doc,
         "The tenant tier mix below is the one assumption with the largest effect on "
         "the CO3 conclusions, so it is worth defending. The distribution is "
         "deliberately skewed towards many small guests, because that is what a "
         "microservice-era platform actually looks like: 900 containers or minimal VMs "
         "with 16 MB working sets, a few hundred mid-sized application guests, and a "
         "small number of genuinely large database and batch guests. A mix weighted "
         "the other way would be easier to make fit and would flatter the design.")

    B.table(doc,
            ["Tier", "Guests", "WSS each", "Frames each", "Tier demand", "GB"],
            [
                ("Web / DNS microservice", "900", "16 MB", "4,096", "3,686,400", "14.06"),
                ("Interactive enterprise app", "240", "48 MB", "12,288", "2,949,120", "11.25"),
                ("Database guest VM", "40", "192 MB", "49,152", "1,966,080", "7.50"),
                ("Batch analytics worker", "20", "160 MB", "40,960", "819,200", "3.12"),
                ("*TOTAL DEMAND  D", "*1,200", "—", "—", "*9,420,800", "*35.94"),
            ],
            "Tenant tier mix and aggregate working-set demand. Available frames after "
            "a 10 % hypervisor and page-cache reserve: 7,549,748 (28.80 GB), so "
            "D/m = 1.248.",
            widths=[1.85, 0.7, 0.8, 0.95, 1.1, 0.65], size=8.6, highlight={4})

    h2(doc, "1.4  Measurable metrics and success criteria")
    para(doc,
         "Each subsystem is judged against a metric that can be computed rather than "
         "argued about. Naming the metric before running the experiment is what stops "
         "the analysis from drifting towards whichever number happened to look best.")

    B.table(doc,
            ["Metric", "Definition", "Target", "Result achieved"],
            [
                ("Page fault ratio", "faults ÷ references", "minimise; monotonic in "
                 "frames", "LRU 55.6 % at 4 frames; monotonic ✔"),
                ("Effective access time", "TLB-weighted mean translation cost",
                 "≪ 1 ms", "184.33 ns ✔"),
                ("Page-table residency", "resident bytes per process",
                 "minimise", "128 KB → 4 KB (32×) ✔"),
                ("Over-commit ratio D/m", "Σ WSS ÷ available frames", "< 1.0",
                 "1.248 → 0.929 after reclaim ✔"),
                ("Mean seek distance", "total head movement ÷ requests",
                 "minimise", "183.08 → 67.75 cylinders ✔"),
                ("Buffer cache hit ratio", "hits ÷ accesses", "maximise",
                 "76.9 % at 1,024 buffers ✔"),
                ("I/O response p99", "99th percentile service latency",
                 "< 10 ms at 90 % load", "23.1 ms — **not met**, see §6.4"),
                ("Allocation success", "requests honoured after churn",
                 "no false rejection", "Indexed/Extent ✔; Contiguous ✘"),
                ("DNS zone integrity", "A records matched by PTR records",
                 "100 %", "15 ↔ 15, all matched ✔"),
                ("Implementation correctness", "validation suite", "100 % pass",
                 "65/65 ✔"),
            ],
            "Measurable metrics, targets and outcomes. Nine of ten targets are met; "
            "the one that is not is analysed honestly in §6.4.",
            widths=[1.35, 1.55, 1.35, 2.25], size=8.2, highlight={6})

    h2(doc, "1.5  Constraints carried from Section D of the brief")
    bullet(doc, "**Functional and operational** — high availability, zero data loss "
                "on crash, strict multi-tenant isolation, sub-10 ms response under "
                "90 %+ RAM load.")
    bullet(doc, "**Memory** — 32 GB RAM, 4 KB pages (or 2 MB hugepages where "
                "justified), 128 MB logical address space, multi-level page tables.")
    bullet(doc, "**Storage** — 500-cylinder unit, 12-request pending queue at peak.")
    bullet(doc, "**Technical and tools** — Linux with root access, BIND 9, Linux "
                "networking tooling, POSIX tracing, Xen or VMware deployment, and "
                "version control.")
    bullet(doc, "**Security and standards** — POSIX file locking, SELinux/AppArmor "
                "policy, ISO/IEC 27001 alignment.")
    para(doc,
         "Section 3.4 maps each of these constraints to the specific artefact that "
         "satisfies it and the evidence file that demonstrates it, including the one "
         "constraint that is only partially satisfied.")
