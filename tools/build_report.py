#!/usr/bin/env python3
"""
CloudMatrix - assignment report builder
=======================================

Assembles the full CSA04 assignment report as a Word document, pulling figures
from results/figures/ and evidence images from screenshots/ so that the
document and the repository can never disagree.

Run:  python tools/build_report.py
Out:  docs/CSA04_OS_Assignment_CO3_CO4_CO5_Tharunkumar_S_192511416.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image as _PILImage

ROOT = Path(__file__).resolve().parent.parent
FIG = ROOT / "results" / "figures"
SHOT = ROOT / "screenshots"
DOCS = ROOT / "docs"
DOCS.mkdir(exist_ok=True)

REPO = "https://github.com/tharuncoder676/OS-ASSIGNMENT"
NAME = "Tharunkumar S"
REGNO = "192511416"

NAVY = RGBColor(0x1F, 0x4E, 0x79)
ACCENT = RGBColor(0xC5, 0x5A, 0x11)
GREY = RGBColor(0x59, 0x59, 0x59)

_fig_n = 0
_tab_n = 0


# --------------------------------------------------------------------------
# document helpers
# --------------------------------------------------------------------------

def setup(doc: Document) -> None:
    s = doc.sections[0]
    s.page_width, s.page_height = Inches(8.27), Inches(11.69)   # A4
    s.left_margin = s.right_margin = Inches(0.85)
    s.top_margin = Inches(0.85)
    s.bottom_margin = Inches(0.8)

    n = doc.styles["Normal"]
    n.font.name = "Calibri"
    n.font.size = Pt(9.8)
    n.paragraph_format.space_after = Pt(4)
    n.paragraph_format.line_spacing = 1.0
    rpr = n.element.get_or_add_rPr()
    rf = OxmlElement("w:rFonts")
    rf.set(qn("w:ascii"), "Calibri")
    rf.set(qn("w:hAnsi"), "Calibri")
    rpr.append(rf)


def page_numbers(doc: Document) -> None:
    for section in doc.sections:
        p = section.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("CSA04 Operating Systems  |  CO3 · CO4 · CO5  |  "
                      f"{NAME} — {REGNO}  |  Page ")
        r.font.size = Pt(8)
        r.font.color.rgb = GREY
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        p._p.append(fld)


def h1(doc, text, num=None):
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(f"{num}. {text}" if num else text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = NAVY
    bar = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "10")
    bot.set(qn("w:color"), "1F4E79")
    bot.set(qn("w:space"), "4")
    bar.append(bot)
    p._p.get_or_add_pPr().append(bar)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11.5)
    r.font.color.rgb = NAVY
    return p


def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = ACCENT
    return p


def para(doc, text, size=9.8, italic=False, align=None):
    p = doc.add_paragraph()
    p.alignment = align or WD_ALIGN_PARAGRAPH.JUSTIFY
    # **bold** segments
    for i, seg in enumerate(text.split("**")):
        if not seg:
            continue
        r = p.add_run(seg)
        r.font.size = Pt(size)
        r.italic = italic
        if i % 2 == 1:
            r.bold = True
    return p


def bullet(doc, text, size=9.8):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.28)
    for i, seg in enumerate(text.split("**")):
        if not seg:
            continue
        r = p.add_run(seg)
        r.font.size = Pt(size)
        if i % 2 == 1:
            r.bold = True
    return p


def mono(doc, text, size=8.1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(4)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(size)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    return p


def shade(cell, hexcolour):
    tcpr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolour)
    tcpr.append(sh)


def table(doc, headers, rows, caption, widths=None, size=8.0, highlight=None):
    global _tab_n
    _tab_n += 1
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(htext)
        r.bold = True
        r.font.size = Pt(size)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(hdr[i], "1F4E79")
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            p.alignment = (WD_ALIGN_PARAGRAPH.LEFT if ci == 0
                           else WD_ALIGN_PARAGRAPH.CENTER)
            txt = str(val)
            bold = txt.startswith("*")
            r = p.add_run(txt.lstrip("*"))
            r.font.size = Pt(size)
            r.bold = bold
            if bold:
                r.font.color.rgb = NAVY
        if highlight and ri in highlight:
            for c in cells:
                shade(c, "FFF2CC")
        elif ri % 2 == 1:
            for c in cells:
                shade(c, "F2F6FA")
    if widths:
        for row in t.rows:
            for i, wd in enumerate(widths):
                row.cells[i].width = Inches(wd)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(f"Table {_tab_n}.  {caption}")
    r.font.size = Pt(8.5)
    r.italic = True
    r.font.color.rgb = GREY
    cap.paragraph_format.space_after = Pt(8)
    return t


# Cap how tall any single image may be, so a long screenshot cannot eat a page.
MAX_H_FIGURE = 2.95      # inches
MAX_H_EVIDENCE = 2.60    # inches


def image(doc, path: Path, caption: str, width=6.4, is_figure=True):
    global _fig_n
    if not path.exists():
        print(f"   !! missing image {path}")
        return
    # Fit the image inside both a width and a height budget.
    with _PILImage.open(path) as im:
        aspect = im.height / im.width
    cap_h = MAX_H_FIGURE if is_figure else MAX_H_EVIDENCE
    if width * aspect > cap_h:
        width = cap_h / aspect
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if is_figure:
        _fig_n += 1
        label = f"Figure {_fig_n}.  "
    else:
        label = ""
    r = cap.add_run(label + caption)
    r.font.size = Pt(8.0)
    r.italic = True
    r.font.color.rgb = GREY
    cap.paragraph_format.space_after = Pt(7)


def evidence(doc, path: Path, caption: str, width=6.4):
    """A screenshot with a heavier 'evidence' caption and a repo pointer."""
    image(doc, path, caption, width)


def pagebreak(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def link(doc, text, url, size=10.5, prefix=""):
    p = doc.add_paragraph()
    if prefix:
        r0 = p.add_run(prefix)
        r0.font.size = Pt(size)
    part = doc.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    col = OxmlElement("w:color"); col.set(qn("w:val"), "0563C1")
    und = OxmlElement("w:u"); und.set(qn("w:val"), "single")
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(size * 2)))
    rpr.append(col); rpr.append(und); rpr.append(sz)
    run.append(rpr)
    t = OxmlElement("w:t"); t.text = text
    run.append(t)
    hl.append(run)
    p._p.append(hl)
    return p
